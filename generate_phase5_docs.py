#!/usr/bin/env python3
"""
generate_phase5_docs.py
Creates:
 - PHASE5_VISUAL_GUIDE.md
 - docs/phase5/screenshots/ (placeholders + README)
 - PHASE5_VISUAL_GUIDE.docx
 - PHASE5_VISUAL_GUIDE.pdf (via pandoc or docx2pdf)

Dependencies:
 - python-docx: pip install python-docx
 - Pillow: pip install Pillow
 - Optional for PDF:
   * pandoc (system package, recommended): brew install pandoc / apt-get install pandoc
   * docx2pdf (Python package): pip install docx2pdf

Usage:
 $ python3 -m venv .venv
 $ source .venv/bin/activate  # On Windows: .venv\\Scripts\\activate
 $ pip install python-docx Pillow
 # For PDF (choose one):
 # - Install pandoc system-wide (recommended)
 # - Or: pip install docx2pdf
 $ python generate_phase5_docs.py
"""
import os
import sys
from pathlib import Path

# Constants for placeholder image generation
PLACEHOLDER_IMG_WIDTH = 1200
PLACEHOLDER_IMG_HEIGHT = 800
PLACEHOLDER_BG_COLOR = (240, 240, 240)
PLACEHOLDER_TEXT_COLOR = (80, 80, 80)
PLACEHOLDER_TEXT_POSITION = (20, 20)

# The full markdown content (expanded version)
MD_CONTENT = r"""
# Phase 5 - TODOS Executive Layout Visual Guide (Validated + Expanded)

_Last updated: 2026-01-13 (Validated & expanded by AI)_

---

## Table of contents
1. Overview
2. Before (Original TODOS Layout) — preserved for reference
3. After (Executive Overview TODOS Layout) — high-level
4. Budget Health Pill — logic & colors
5. Collapsible Sections Behavior & Accessibility
6. Modes: TODOS (Portfolio) and Single-Project
7. Component Hierarchy (full module list)
8. Module & Component Specifications (API, props, data contracts)
9. Design Tokens & Visual Assets
10. Code Structure & Key Files
11. Testing & QA Plan (Visual, Functional, Data Integrity, E2E, Performance)
12. Screenshots Required & Naming Conventions
13. Implementation Notes, Acceptance Criteria & Rollback Plan
14. Change Log

---

## 1. Overview
This guide documents the Phase 5 transformation of the TODOS view from a dense, information-heavy layout to an executive-friendly overview. The emphasis is on immediate clarity, progressive disclosure, visible budget health, and preserving all existing features and behaviors in single-project mode.

---

## 2. Before (Original TODOS Layout)
(Reference — unchanged)
- Header: Gestión de Pronóstico + actions
- Baseline Status Panel (shown for all projects)
- KPI Summary Cards (6 cards)
- Budget Simulation KPIs (when enabled)
- Real Annual Budget KPIs
- Budget & Simulation Panel (collapsible)
- Portfolio Summary View
- Charts Panel
- Rubros Table
- Forecast Grid (12-month)
- Charts and analytics
- Issues: Too much above the fold, unclear hierarchy, charts buried, no way to hide heavy sections.

---

## 3. After (Executive Overview TODOS Layout)
(High level)
- Header: Gestión de Pronóstico + actions
- **Executive KPI Summary Bar** (top of page; shows Budget Health pill + top-line numbers)
- **Charts Panel** prominent and near top (Monthly Trend, By Category, Cumulative)
- Collapsible sections (default closed): Portfolio Summary, Forecast Grid (12 months), Budget Simulator, Desglose Mensual vs Presupuesto — each section preserves original content but is hidden behind progressive disclosure
- Single project mode remains unchanged.

---

## 4. Budget Health Pill — Status Logic & Colors

### Status logic (unchanged)
- **🟢 EN META**  
  - Conditions: `consumption <= 90% AND forecast <= budget`  
  - Color: Green background, green text  
  - Meaning: On track.
- **🟡 EN RIESGO**  
  - Conditions: `consumption > 90% AND forecast <= budget`  
  - Color: Yellow background, yellow text  
  - Meaning: High consumption but still under budget.
- **🔴 SOBRE PRESUPUESTO**  
  - Conditions: `forecast > budget OR consumption > 100%`  
  - Color: Red background, red text  
  - Meaning: Over budget/overspent.
- **⚪ SIN PRESUPUESTO**  
  - Conditions: `no monthly budget configured`  
  - Color: Gray background, gray text  
  - Meaning: Budget not set up yet.

### Suggested color tokens (HEX)
- `--color-budget-green: #4ac795`  *(Ikusi green as requested for table header color)*  
- `--color-budget-green-dark: #3dbf82`  
- `--color-budget-yellow: #F2C94C`  
- `--color-budget-red: #FF6B6B`  
- `--color-budget-gray: #BDBDBD`  
> Note: adjust for WCAG contrast as needed. See Accessibility section.

---

## 5. Collapsible Sections Behavior & Accessibility

**Behavior**
- Default: *closed* for all four main sections.
- Interaction: click or press Enter/Space on the header/chevron toggles open/closed.
- Smooth animation and progressive disclosure; content loads when expanded (lazy load to reduce initial payload).

**Accessibility**
- Use native `<button>` for triggers or `role="button"` with `tabindex="0"`.
- ARIA attributes:
  - `aria-expanded="false|true"` on the trigger element.
  - `aria-controls="section-id"` linking to the collapsible content.
  - Use `id` on collapsible content. Example: `id="portfolio-summary-content"`.
- Keyboard:
  - `Tab` to focus trigger, `Enter` / `Space` toggles.
  - If sections must be navigable quickly, consider `Alt + Arrow` patterns for accordion navigation.
- Focus management:
  - When opened, focus stays on the trigger (do not auto-focus into big content that disorients users).
  - For screen reader users, announce "expanded" state via `aria-live="polite"` for the KPI pill updates if the change is dynamic.

---

## 6. Modes: TODOS (Portfolio) and Single-Project

**TODOS (isPortfolioView === true)**
- Baseline panel hidden.
- KPI cards hidden.
- Executive KPI Summary Bar shown at top (ForecastSummaryBar).
- Charts Panel shown (ForecastChartsPanel).
- Four collapsible wrapper sections introduced:
  - `Resumen de todos los proyectos` → `PortfolioSummaryView`
  - `Cuadrícula de Pronóstico (12 meses)` → `ForecastRubrosTable`
  - `Simulador de Presupuesto` → Annual/Monthly budget editor + simulator
  - `Desglose Mensual vs Presupuesto` → full Forecast Grid
- Default: collapsed (progressive disclosure).

**Single Project (isPortfolioView === false)**
- Remains unchanged. Baseline visible, KPI cards visible, original Forecast Grid and charts available.

---

## 7. Component Hierarchy (full module list)

**Top-level**
- `SDMTForecast` (root)
  - `Header + Actions`
  - `DataHealthPanel` (dev only)
  - `ForecastSummaryBar` *(NEW position for TODOS)*
    - `BudgetHealthPill` *(NEW component)*
    - `TotalsRow` (Budget/Forecast/Actual/%Consumption/Variance)
  - `ForecastChartsPanel` *(NEW position for TODOS)*
    - `MonthlyTrendChart`, `ByCategoryChart`, `CumulativeChart`
  - `CollapsibleSection` wrappers (4):
    - `PortfolioSummaryView`
    - `ForecastRubrosTable`
    - `BudgetSimulatorPanel` (Annual + Monthly)
    - `ForecastGridDetailed`
  - `BaselineStatusPanel` (single project only)
  - `KpiSummaryCards` (single project only)
  - `ChartInsightsPanel` (single project only)

---

## 8. Module & Component Specifications

> Each component includes props and data contract expectations. Keep props minimal and typed (TypeScript).

### ForecastSummaryBar.tsx
**Responsibility:** Display executive KPI summary across portfolio with Budget Health pill.

**Props**
```ts
interface ForecastSummaryBarProps {
  totalBudget: number;         // cents or decimal (document unit)
  totalForecast: number;
  totalActual: number;
  percentConsumption: number;  // 0-100
  variance: number;            // totalForecast - totalBudget
  lastUpdated?: string;        // ISO timestamp
  budgetHealth?: 'EN_META' | 'EN_RIESGO' | 'SOBRE_PRESUPUESTO' | 'SIN_PRESUPUESTO';
}
```

**Behavior**
- Derive `budgetHealth` from the values if not passed.
- Render `BudgetHealthPill` with appropriate color and label.
- Expose `data-testid` attributes for e2e tests.

---

### BudgetHealthPill.tsx

**Props**
```ts
interface BudgetHealthPillProps {
  status: 'EN_META' | 'EN_RIESGO' | 'SOBRE_PRESUPUESTO' | 'SIN_PRESUPUESTO';
  tooltipText?: string;
}
```

**Accessibility**
- `role="status"` and `aria-live="polite"` for dynamic updates.
- High contrast text vs background (WCAG AA minimum).

---

### Collapsible (generic)

**Props**
```ts
interface CollapsibleProps {
  id: string;
  headerLabel: string;
  defaultOpen?: boolean;
  onToggle?: (open: boolean) => void;
  children: React.ReactNode;
}
```

**Accessibility**
- Trigger: `<button aria-expanded aria-controls>`
- Content: `<div id={id} role="region" aria-labelledby={triggerId}>`

---

### ForecastRubrosTable / ForecastGridDetailed

**Requirements**
- Table header color: Ikusi green `--color-ikusi-green: #4ac795` for header backgrounds.
- Support large datasets with virtualization (react-window / react-virtualized).
- Inline editing cells must persist edits with optimistic UI and rollback on server error.

---

## 9. Design Tokens & Visual Assets

**Color palette**
- `--color-ikusi-green: #4ac795` (table header & accents)
- `--color-budget-green: #4ac795`
- `--color-budget-yellow: #F2C94C`
- `--color-budget-red: #FF6B6B`
- `--color-budget-gray: #BDBDBD`
- `--color-card-bg: #FFFFFF`
- `--color-muted-text: #6B6B6B`

**Typography**
- Headline: Inter 600 (or system sans), 20-24px
- Body: Inter 400, 14-16px
- Small / secondary: Inter 400, 12px

**Spacing**
- 8px base spacing (grid) — use multiples for layout.

**Icons & Logos**
- Use SVG icons (inline) for KPI pill, chevrons, status badges.
- Ensure exported SVGs have `role="img"` and `aria-label`.

---

## 10. Code Structure & Key Files

- `src/views/SDMTForecast.tsx` (root)
- `src/components/ForecastSummaryBar/ForecastSummaryBar.tsx`
- `src/components/BudgetHealthPill/BudgetHealthPill.tsx`
- `src/components/ForecastChartsPanel/ForecastChartsPanel.tsx`
- `src/components/Collapsible/Collapsible.tsx`
- `src/components/ForecastRubrosTable/ForecastRubrosTable.tsx`
- `src/components/ForecastGridDetailed/ForecastGridDetailed.tsx`
- `src/styles/tokens.css` (or tokens.ts)
- `src/tests/visual/phase5/` (visual regression tests)
- `src/tests/unit/ForecastSummaryBar.test.tsx` (unit)
- `src/e2e/phase5.spec.ts` (cypress / playwright)

**Sample code snippet (Budget Health logic):**
```ts
export function getBudgetHealthStatus({ consumption, forecast, budget }) {
  if (!budget) return { label: 'SIN PRESUPUESTO', status: 'SIN_PRESUPUESTO', color: '--color-budget-gray' };
  if (forecast > budget || consumption > 100) return { label: 'SOBRE PRESUPUESTO', status: 'SOBRE_PRESUPUESTO', color: '--color-budget-red' };
  if (consumption > 90) return { label: 'EN RIESGO', status: 'EN_RIESGO', color: '--color-budget-yellow' };
  return { label: 'EN META', status: 'EN_META', color: '--color-budget-green' };
}
```

---

## 11. Testing & QA Plan

### Visual Testing
- Verify the Executive KPI bar is visible above the fold in TODOS view.
- Confirm Budget Health pill displays correct color for each status.
- Charts panel visible and interactive (hover tooltips, legend toggles).
- Collapsible sections — all 4 collapsed by default; expand/collapse works and content loads.

### Functional Testing
- Unit tests for `getBudgetHealthStatus` (all branches).
- Unit tests for `ForecastSummaryBar` rendering with various props.
- Integration tests verifying collapsed sections lazy-load data.
- E2E tests:
  - Navigate to TODOS view, validate above the fold items.
  - Toggle sections and validate content loads with correct data.
  - Test inline edits in `ForecastGridDetailed` and persistence/rollback on error.

### Data Integrity
- Ensure totals match the underlying dataset.
- Confirm no rounding errors or currency formatting issues.
- Validate budget simulation results match server-calculated values.

### Performance
- Page initial load time < 3s on 3G simulated network.
- Collapsible lazy-loading should reduce payload for TODOS view.
- Virtualization for tables when > 200 rows.

### Regression
- Single-project mode unchanged. Run smoke tests to confirm baseline panel/KPI cards exist.

---

## 12. Screenshots Required & Naming Conventions

**TODOS View**
1. `phase5_todos_above_fold.png` — Executive KPI bar + Charts panel
2. `phase5_budget_pill_en_meta.png` — Budget Health pill (EN META)
3. `phase5_budget_pill_en_riesgo.png` — Budget Health pill (EN RIESGO)
4. `phase5_budget_pill_sobre_presupuesto.png` — Budget Health pill (SOBRE PRESUPUESTO)
5. `phase5_collapsed_sections.png` — All 4 sections collapsed
6. `phase5_expanded_portfolio_summary.png` — Example expanded section
7. `phase5_full_page_scroll.png` — Full TODOS layout

**Single Project**
8. `phase5_single_above_fold.png` — Baseline panel + KPI cards
9. `phase5_single_full_layout.png` — Full single-project layout

**Comparison**
10. `phase5_before_after_todos.png` — side-by-side comparison

**Storage**
- Place images under: `docs/phase5/screenshots/` in the repo.

---

## 13. Implementation Notes, Acceptance Criteria & Rollback Plan

**Acceptance criteria**
- Executive KPI Summary Bar is displayed on TODOS by default and reflects calculated totals.
- Budget Health pill displays correct color and label for all four statuses.
- Collapsible wrappers exist and lazy-load their content when opened.
- Single-project mode is unchanged.
- All tests in `src/tests/*` pass in CI (unit + e2e + visual regression).

**Rollback**
- Keep the feature behind a feature flag: `FEATURE_TODOS_EXECUTIVE_LAYOUT`.
- If any production issue arises, flip the flag off to restore previous TODOS layout instantaneously.

---

## 14. Change Log

- `2026-01-13` — Validated original file; expanded with module APIs, accessibility, design tokens, full QA plan, screenshot conventions, and implementation notes.

---

## Appendix: Quick dev checklist (copy to PR template)

- [ ] Add `ForecastSummaryBar` component and `BudgetHealthPill`.
- [ ] Add tokens file and include Ikusi green for table header.
- [ ] Implement collapsible wrapper with aria attributes.
- [ ] Add unit tests and e2e tests for TODOS behaviors.
- [ ] Add visual regression tests and collect baseline screenshots.
- [ ] Add `FEATURE_TODOS_EXECUTIVE_LAYOUT` feature flag with default enabled in staging.
- [ ] Update PR description with before/after screenshots.
"""

def write_file(path: Path, content: str):
    """Write content to file, creating parent directories as needed."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    print(f"✅ Wrote {path}")

def create_screenshot_placeholders(base_dir: Path):
    """Create placeholder PNG files for all required screenshots."""
    base_dir.mkdir(parents=True, exist_ok=True)
    filenames = [
        "phase5_todos_above_fold.png",
        "phase5_budget_pill_en_meta.png",
        "phase5_budget_pill_en_riesgo.png",
        "phase5_budget_pill_sobre_presupuesto.png",
        "phase5_collapsed_sections.png",
        "phase5_expanded_portfolio_summary.png",
        "phase5_full_page_scroll.png",
        "phase5_single_above_fold.png",
        "phase5_single_full_layout.png",
        "phase5_before_after_todos.png",
    ]
    for fn in filenames:
        p = base_dir / fn
        if not p.exists():
            # Create a placeholder PNG with text
            try:
                from PIL import Image, ImageDraw, ImageFont
                img = Image.new("RGB", (PLACEHOLDER_IMG_WIDTH, PLACEHOLDER_IMG_HEIGHT), color=PLACEHOLDER_BG_COLOR)
                d = ImageDraw.Draw(img)
                text = f"Placeholder: {fn}"
                # Use default font
                d.text(PLACEHOLDER_TEXT_POSITION, text, fill=PLACEHOLDER_TEXT_COLOR)
                img.save(p, "PNG")
                print(f"✅ Created placeholder {fn}")
            except ImportError:
                # Fallback: create empty file if PIL is not available
                p.write_bytes(b"")
                print(f"⚠️  Created empty placeholder {fn} (Pillow not installed)")
            except Exception as e:
                # Fallback: create empty file on any error
                p.write_bytes(b"")
                print(f"⚠️  Created empty placeholder {fn} (Error: {e})")

def write_screenshots_readme(base_dir: Path):
    """Create README.md in screenshots directory with guidelines."""
    content = """# PHASE5 Screenshots

Place the required screenshots here following the naming convention.

## Required files

1. `phase5_todos_above_fold.png` - Executive KPI bar + Charts panel visible above the fold
2. `phase5_budget_pill_en_meta.png` - Budget Health pill showing "EN META" status (green)
3. `phase5_budget_pill_en_riesgo.png` - Budget Health pill showing "EN RIESGO" status (yellow)
4. `phase5_budget_pill_sobre_presupuesto.png` - Budget Health pill showing "SOBRE PRESUPUESTO" status (red)
5. `phase5_collapsed_sections.png` - All 4 collapsible sections in collapsed state
6. `phase5_expanded_portfolio_summary.png` - Example of an expanded collapsible section
7. `phase5_full_page_scroll.png` - Full TODOS layout from top to bottom
8. `phase5_single_above_fold.png` - Single project view showing baseline panel + KPI cards
9. `phase5_single_full_layout.png` - Full single-project layout
10. `phase5_before_after_todos.png` - Side-by-side comparison of before/after layouts

## Guidelines

- **Resolution**: Use 1200x800 PNG format with high quality
- **Accessibility**: Ensure accessibility states are visible in the pill screenshots
- **Consistency**: Use the same project/data across screenshots for consistency
- **Clarity**: Ensure text is readable and UI elements are clearly visible
- **Context**: Include enough context to understand the feature being demonstrated
"""
    write_file(base_dir / "README.md", content)

def create_docx(md_text: str, out_path: Path):
    """Convert markdown to DOCX using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        print("❌ python-docx not installed. Please run: pip install python-docx")
        raise e
    
    # Simple markdown -> docx conversion
    doc = Document()
    
    in_code_block = False
    code_block_lines = []
    
    for line in md_text.splitlines():
        # Handle code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End of code block
                code_text = "\n".join(code_block_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                # Make code blocks slightly different
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                code_block_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            continue
        
        # Handle headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        else:
            # Regular paragraph
            if line.strip():  # Skip empty lines
                doc.add_paragraph(line)
    
    doc.save(out_path)
    print(f"✅ Wrote DOCX to {out_path}")

def convert_docx_to_pdf(docx_path: Path, pdf_path: Path):
    """Convert DOCX to PDF using pandoc or docx2pdf."""
    import subprocess
    import shutil
    
    # Try pandoc first (preferred)
    if shutil.which("pandoc"):
        cmd = ["pandoc", str(docx_path), "-o", str(pdf_path)]
        try:
            subprocess.check_call(cmd, stderr=subprocess.PIPE)
            print(f"✅ Generated PDF via pandoc: {pdf_path}")
            return True
        except subprocess.CalledProcessError as e:
            print(f"⚠️  Pandoc conversion failed: {e}")
    else:
        print("ℹ️  Pandoc not found in PATH")
    
    # Fallback to docx2pdf
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"✅ Generated PDF via docx2pdf: {pdf_path}")
        return True
    except ImportError:
        print("⚠️  docx2pdf not installed. Please run: pip install docx2pdf")
        print("    Or install pandoc system-wide (recommended)")
        return False
    except Exception as e:
        print(f"⚠️  docx2pdf conversion failed: {e}")
        return False

def print_dependency_instructions():
    """Print dependency installation instructions."""
    print("\n" + "="*70)
    print("📦 DEPENDENCY INSTALLATION INSTRUCTIONS")
    print("="*70)
    print("\nRequired dependencies:")
    print("  - python-docx: pip install python-docx")
    print("  - Pillow: pip install Pillow")
    print("\nOptional dependencies for PDF generation (choose one):")
    print("  - pandoc (recommended, system package):")
    print("    * macOS: brew install pandoc")
    print("    * Ubuntu/Debian: sudo apt-get install pandoc")
    print("    * Windows: Download from https://pandoc.org/installing.html")
    print("  - docx2pdf (Python package, Windows or with wine on Linux):")
    print("    * pip install docx2pdf")
    print("\nTo install all Python dependencies:")
    print("  pip install python-docx Pillow")
    print("="*70 + "\n")

def print_summary(files_created, repo_root):
    """Print summary of created files and next steps."""
    print("\n" + "="*70)
    print("✨ SUMMARY")
    print("="*70)
    print("\nFiles created:")
    for f in files_created:
        rel_path = f.relative_to(repo_root) if f.is_absolute() else f
        print(f"  ✅ {rel_path}")
    
    print("\n" + "="*70)
    print("🚀 NEXT STEPS - Git Commands")
    print("="*70)
    print("\nTo commit and push these changes on a new branch:")
    print("\n  git checkout -b feature/phase5-docs")
    print("  git add PHASE5_VISUAL_GUIDE.md PHASE5_VISUAL_GUIDE.docx docs/phase5/")
    
    # Check if PDF was created
    pdf_path = repo_root / "PHASE5_VISUAL_GUIDE.pdf"
    if pdf_path.exists():
        print("  git add PHASE5_VISUAL_GUIDE.pdf")
    else:
        print("  # Note: PDF was not generated. Install pandoc or docx2pdf and re-run")
    
    print('  git commit -m "PHASE5: Add validated visual guide + DOCX/PDF and screenshots placeholders"')
    print("  git push --set-upstream origin feature/phase5-docs")
    print("\n" + "="*70 + "\n")

def main():
    """Main execution function."""
    print("\n" + "="*70)
    print("🎯 Phase 5 Documentation Generator")
    print("="*70 + "\n")
    
    print_dependency_instructions()
    
    repo_root = Path.cwd()
    files_created = []
    
    # 1. Create/overwrite PHASE5_VISUAL_GUIDE.md
    print("📝 Creating PHASE5_VISUAL_GUIDE.md...")
    md_path = repo_root / "PHASE5_VISUAL_GUIDE.md"
    write_file(md_path, MD_CONTENT)
    files_created.append(md_path)
    
    # 2. Create screenshots directory and placeholders
    print("\n📸 Creating screenshots directory and placeholders...")
    screenshots_dir = repo_root / "docs" / "phase5" / "screenshots"
    create_screenshot_placeholders(screenshots_dir)
    write_screenshots_readme(screenshots_dir)
    files_created.append(screenshots_dir)
    
    # 3. Create DOCX
    print("\n📄 Creating PHASE5_VISUAL_GUIDE.docx...")
    docx_path = repo_root / "PHASE5_VISUAL_GUIDE.docx"
    try:
        create_docx(MD_CONTENT, docx_path)
        files_created.append(docx_path)
    except Exception as e:
        print(f"❌ Failed to create DOCX: {e}")
        print("Please install python-docx: pip install python-docx")
        sys.exit(1)
    
    # 4. Create PDF
    print("\n📑 Creating PHASE5_VISUAL_GUIDE.pdf...")
    pdf_path = repo_root / "PHASE5_VISUAL_GUIDE.pdf"
    pdf_success = convert_docx_to_pdf(docx_path, pdf_path)
    if pdf_success and pdf_path.exists():
        files_created.append(pdf_path)
    else:
        print("\n⚠️  PDF was not generated.")
        print("To generate PDF, install one of the following:")
        print("  1. pandoc (recommended): brew install pandoc / apt-get install pandoc")
        print("  2. docx2pdf: pip install docx2pdf")
        print("Then re-run this script.")
    
    # 5. Print summary
    print_summary(files_created, repo_root)

if __name__ == "__main__":
    main()
